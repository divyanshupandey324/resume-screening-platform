import pdfplumber
import zipfile
import re

def extract_resume_text(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    return text

def extract_docx_text(docx_path):
    try:
        import docx
        doc = docx.Document(docx_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)
    except Exception as e:
        print("DOCX parsing error:", e)
        # Fallback reading zip archive
        return extract_xml_zip_text(docx_path, 'word/document.xml')

def extract_pptx_text(pptx_path):
    text = []
    try:
        with zipfile.ZipFile(pptx_path, 'r') as zip_ref:
            # Sort files to read in slide order
            slide_files = sorted([f for f in zip_ref.namelist() if re.match(r'ppt/slides/slide\d+\.xml', f)])
            for slide_file in slide_files:
                slide_content = zip_ref.read(slide_file).decode('utf-8', errors='ignore')
                # Find all text tags <a:t>...</a:t>
                matches = re.findall(r'<a:t[^>]*>(.*?)</a:t>', slide_content)
                text.extend(matches)
    except Exception as e:
        print("PPTX parsing error:", e)
    return "\n".join(text)

def extract_xml_zip_text(zip_path, xml_filename):
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            xml_content = zip_ref.read(xml_filename).decode('utf-8', errors='ignore')
            # Extract all tags of text
            matches = re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml_content)
            return "\n".join(matches)
    except Exception as e:
        print(f"Zip XML extraction error for {xml_filename}:", e)
        return ""

def extract_any_document(file_path):
    ext = file_path.lower().split('.')[-1]
    if ext == 'pdf':
        return extract_resume_text(file_path)
    elif ext == 'docx':
        return extract_docx_text(file_path)
    elif ext in ['ppt', 'pptx']:
        return extract_pptx_text(file_path)
    else:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            return ""